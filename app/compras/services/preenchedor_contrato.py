import io
import re
import logging
from datetime import date
from decimal import Decimal
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from compras.models import Contrato, VinculoOrganizacao, Organizacao, Compra, Empenho


logger = logging.getLogger(__name__)


class PreenchedorContratoService():

   
    @staticmethod
    def localizar_secao(text, current_state):
        """
        Locate the section (context state) based on the text.
        Returns the new state or current_state.
        """
        text_upper = text.upper()
        
        # Check for Representative context first as it is very specific
        if "REPRESENTANTE LEGAL DO CONTRATANTE" in text_upper or "REPRESENTANTE LEGAL DA CONTRATANTE" in text_upper:
            return 'REPRESENTANTE_CONTRATANTE'
        if "REPRESENTANTE LEGAL DO CONTRATADO" in text_upper or "REPRESENTANTE LEGAL DA CONTRATADA" in text_upper or "REPRESENTANTE LEGAL DO CONTRATADA" in text_upper:
            return 'REPRESENTANTE_CONTRATADA'
            
        # Check for general Contratante/Contratada context
        if "DADOS DA CONTRATANTE" in text_upper:
            return 'CONTRATANTE'
        if "DADOS DA CONTRATADA" in text_upper or "DADOS DO CONTRATADO" in text_upper:
            return 'CONTRATADA'
            
        return current_state

    @classmethod
    def substituir_texto(cls, paragraph, target, replacement):
        """
        Substitute target text with replacement inside a python-docx paragraph,
        preserving formatting (runs) as much as possible.
        """
        if not target or target not in paragraph.text:
            return False
            
        if target == replacement:
            return False
            
        # Try simple run replacement
        replaced = False
        for run in paragraph.runs:
            if target in run.text:
                run.text = run.text.replace(target, replacement)
                replaced = True
                
        if replaced: #and target not in paragraph.text:
            return True
            
        # Reconstruct runs for split targets
        text = paragraph.text
        idx = text.find(target)
        max_loops = text.count(target)
        loop_count = 0
        
        while idx != -1 and loop_count < max_loops:
            loop_count += 1
            start_run_idx = -1
            end_run_idx = -1
            current_len = 0
            
            for i, run in enumerate(paragraph.runs):
                run_len = len(run.text)
                if start_run_idx == -1 and current_len + run_len > idx:
                    start_run_idx = i
                if end_run_idx == -1 and current_len + run_len >= idx + len(target):
                    end_run_idx = i
                    break
                current_len += run_len
                
            if start_run_idx != -1 and end_run_idx != -1:
                combined_text = ""
                for i in range(start_run_idx, end_run_idx + 1):
                    combined_text += paragraph.runs[i].text
                    
                local_idx = combined_text.find(target)
                if local_idx != -1:
                    new_combined = combined_text[:local_idx] + replacement + combined_text[local_idx + len(target):]
                    paragraph.runs[start_run_idx].text = new_combined
                    for i in range(start_run_idx + 1, end_run_idx + 1):
                        paragraph.runs[i].text = ""
            else:
                break
            
            text = paragraph.text
            idx = text.find(target)
            
        return True

    @classmethod
    def preencher_labels_contextuais(cls, paragraph, context_state, data):
        """
        Fill contextual labels based on the current context state.
        Handles both literal string replacements and signature block replacements.
        """
        text = paragraph.text
        
        # 1. Signature fields (Nome) under representative contexts
        
        # 2. General contextual replacements
        mappings = {}
        if context_state == 'PREAMBULO':
            mappings = {
                "[UNIDADE]": data.get("contratante_nome", ""),
                "[NOME DA EMPRESA]": data.get("contratada_nome", ""),
            }
        elif context_state == 'CONTRATANTE':
            mappings = {
                "[UNIDADE]": data.get("contratante_nome_fantasia", ""),
                "[CNPJ nº]": data.get("contratante_cnpj", ""),
                
                "[endereço completo]": (
                    f"{data.get('contratante_endereco', '')} - "
                    f"{data.get('contratante_estado', '')}"
                 ),

                "[cargo da autoridade competente e nome]": data.get("contratante_resp_completo", ""),
            }
        elif context_state == 'CONTRATADA':
            mappings = {
                "[NOME DA EMPRESA]": data.get("contratada_nome", ""),
                "[CNPJ nº]": data.get("contratada_cnpj", ""),
                "[endereço completo]": data.get("contratada_endereco_completo", ""),
            }

        for target, replacement in mappings.items():
            cls.substituir_texto(paragraph, target, replacement)

    @classmethod
    def _preencher_assinatura(cls, paragraph, valor):
        """Fill signature line with the name by replacing underscores or placeholders."""
        if not valor:  # Se não há valor, não faz nada
            return
        
        # Tentar substituir placeholder primeiro {{ nome }}
        if "{{ nome }}" in paragraph.text:
            cls.substituir_texto(paragraph, "{{ nome }}", valor)
            return
        
        # Procura e substitui underscores nos runs (preservando formatação)
        for run in paragraph.runs:
            if re.search(r'_{3,}', run.text):
                run.text = re.sub(r'_{3,}', valor, run.text, count=1)
                return
        
        # Se não achou nos runs, substitui no parágrafo todo
        if re.search(r'_{3,}', paragraph.text):
            paragraph.text = re.sub(r'_{3,}', valor, paragraph.text, count=1)

    @classmethod
    def processar_paragrafo_geral(cls, paragraph, data):
        """Process general placeholders that are independent of context state."""
        
        # Remover linha "Plano Interno:" se existir
        if "Plano Interno:" in paragraph.text:
            p = paragraph._element
            p.getparent().remove(p)
            return
            
        if "ANEXO III – MINUTA DE TERMO DE CONTRATO" in paragraph.text:
            novo_texto = paragraph.text.replace(
                "ANEXO III – MINUTA DE TERMO DE CONTRATO",
                "TERMO DE CONTRATO"
            )

            if paragraph.runs:
                paragraph.runs[0].text = novo_texto

                for run in paragraph.runs[1:]:
                    run.text = ""
      
      
        if "Processo SEI" in paragraph.text:
            m = re.search(
                r'Processo SEI nº ([0-9NNAA./_-]+)',
                paragraph.text
            )

            if m:
                cls.substituir_texto(
                    paragraph,
                    m.group(1),
                    data.get('numero_sei', '')
                )

        # Número do contrato
        
        if "Contrato nº" in paragraph.text:
            novo_texto = (
                f"Contrato nº {data.get('contrato_numero', '')} "
                f"{data.get('contratante_nome_fantasia', '')}"
            )

            if paragraph.runs:
                paragraph.runs[0].text = novo_texto

                for run in paragraph.runs[1:]:
                    run.text = ""

        # Contratada
        substituicoes = {
            "[CNPJ nº]": data.get("contratada_cnpj", ""),
            "[endereço completo]": data.get("contratada_endereco_completo", ""),
        }

        for placeholder, valor in substituicoes.items():
            if placeholder in paragraph.text:
                cls.substituir_texto(paragraph, placeholder, valor)


        # Objeto
        cls.substituir_texto(paragraph, "[DESCRIÇÃO SUCINTA DO OBJETO]", data.get("objeto", ""))
        
        # Data do orçamento estimado (DEVE VIR ANTES DAS SUBSTITUIÇÕES GENÉRICAS DE DD/MM/AAAA)
        if "data do orçamento estimado, em" in paragraph.text.lower():
            cls.substituir_texto(paragraph, "DD/MM/AAAA", data.get("data_estimativa", ""))
        
        # Data por extenso (DD de MMM de AAAA)
        cls.substituir_texto(paragraph, "DD", data.get("data_dd", ""))
        cls.substituir_texto(paragraph, "MMM", data.get("data_mmm", "").upper())
        cls.substituir_texto(paragraph, "AAAA", data.get("data_aaaa", ""))
        
        # Proposta comercial
        cls.substituir_texto(paragraph, "[NN/NN/NNNN]", data.get("data_proposta", ""))
        
        # Valor do contrato (R$ e extenso)
        if "valor total" in paragraph.text.lower() or "valor da contratação" in paragraph.text.lower():
            replacement = (
                f"{data.get('valor_formatado', '')} "
                f"({data.get('valor_extenso', '')})"
            )

            cls.substituir_texto(
                paragraph,
                "R$.......... (.....)",
                replacement
            )

            cls.substituir_texto(
                paragraph,
                "R$ __________ (__________)",
                replacement
            )
            
        # Garantia contratual
        if "modalidade" in paragraph.text.lower():
            cls.substituir_texto(
                paragraph,
                "__________________",
                f"{data.get('garantia_modalidade', '')}"
            )

        if "valor de r$" in paragraph.text.lower():
            cls.substituir_texto(
                paragraph,
                "R$ _______________",
                f"{data.get('garantia_valor', '')}"
            )

                    
        
        # Preencher campos de empenho (placeholders individuais)
        print("PARÁGRAFO:", repr(paragraph.text))

        for i, run in enumerate(paragraph.runs):
            print(f"RUN {i} =", repr(run.text))
        
        cls.substituir_texto(paragraph, "Gestão/Unidade:", f"Gestão/Unidade: {data.get('empenho_unidade', '')}")
        
        cls.substituir_texto(paragraph, "Fonte de Recursos:", f"Fonte de Recursos: {data.get('empenho_fonte_recurso', '')}")
        cls.substituir_texto(paragraph, "Programa de Trabalho:", f"Programa de Trabalho: {data.get('empenho_programa_trabalho', '')}")
        cls.substituir_texto(paragraph, "Elemento de Despesa:", f"Elemento de Despesa: {data.get('empenho_elemento', '')}")
        cls.substituir_texto(paragraph, "Nota de Empenho:", f"Número do Empenho: {data.get('empenho_numero', '')}")

        
        # Local e data da assinatura
        cls.substituir_texto(paragraph, "[Local]", data.get("contratante_cidade", ""))
        cls.substituir_texto(paragraph, "[dia]", data.get("assinatura_dia", ""))
        cls.substituir_texto(paragraph, "[mês]", data.get("assinatura_mes", ""))
        cls.substituir_texto(paragraph, "[ano]", data.get("assinatura_ano", ""))
        
        # Placeholders específicos de assinatura
        if "Representante legal do CONTRATANTE" in paragraph.text:
            nome_p = paragraph.insert_paragraph_before()
            nome_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = nome_p.add_run(
                data.get("contratante_resp_nome", "")
            )
            run.bold = True
            run.font.size = Pt(11)
                    

        if "Representante legal do CONTRATADO" in paragraph.text:
            nome_p = paragraph.insert_paragraph_before()
            nome_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = nome_p.add_run(
                data.get("contratada_resp_nome", "")
            )
            run.bold = True
            run.font.size = Pt(11)
        


    @classmethod
    def fill_docx(cls, docx_file, contrato):
        """
        Main entry point to fill the docx contract template.
        Reads docx_file, extracts data from contrato, replaces placeholders,
        and returns the filled DOCX file as an io.BytesIO buffer and a filename.
        """
        

        # 1. Fetch related objects
        compra = contrato.compra
        contratante = contrato.contratante
        contratada = contrato.contratada
        empenho = None
        
        # Try to get empenho if it exists
        try:
            empenho = contrato.empenho
        except:
            empenho = None
        
        # 2. Get representatives
        v_contratante = VinculoOrganizacao.objects.filter(
            organizacao=contratante, responsavel_assinatura=True, ativo=True
        ).select_related('pessoa').first()
        
        if contrato.representante_contratada:
            v_contratada = contrato.representante_contratada
        else:
            v_contratada = VinculoOrganizacao.objects.filter(
                organizacao=contratada, responsavel_assinatura=True, ativo=True
            ).select_related('pessoa').first()
        
        # 3. Build data dictionary
        data = {
            "numero_sei": compra.numero_sei if compra else "",
            "contrato_numero": contrato.numero,
            "contratante_nome": contratante.nome,
            "contratante_nome_fantasia": contratante.nome_fantasia or contratante.nome,
            "contratante_cnpj": contratante.cnpj,
            "contratante_endereco": contratante.endereco or "",
            "contratante_cidade": contratante.cidade or "",
            
            "contratada_nome": contratada.nome,
            "contratada_cnpj": contratada.cnpj,
            "contratada_endereco": contratada.endereco or "",
            "contratada_cidade": contratada.cidade or "",
            "contratada_endereco_completo": f"{contratada.endereco or ''}, {contratada.cidade or ''} - {contratada.estado or ''} " .strip(", "),
            
            "objeto": compra.objeto if compra else "",
        }
        
        # Representatives data
        if v_contratante:
            data["contratante_resp_nome"] = v_contratante.pessoa.nome
        else:
            data["contratante_resp_nome"] = ""
        if v_contratada:
            data["contratada_resp_nome"] = v_contratada.pessoa.nome
        else:
            data["contratada_resp_nome"] = ""

            
        # Empenho data
        if empenho:
            data["empenho_unidade"] = empenho.unidade or ""
            data["empenho_fonte_recurso"] = empenho.fonte_recurso or ""
            data["empenho_programa_trabalho"] = "122 - Administração Geral"
            data["empenho_elemento"] = empenho.elemento or ""
            data["empenho_numero"] = empenho.numero or ""
        else:
            data["empenho_unidade"] = ""
            data["empenho_fonte_recurso"] = ""
            data["empenho_programa_trabalho"] = "122 - Administração Geral"
            data["empenho_elemento"] = ""
            data["empenho_numero"] = ""
        # Contract Dates
        if contrato.data:
            dd, mmm, aaaa = contrato.data_por_extenso
            data["data_dd"] = dd
            data["data_mmm"] = mmm
            data["data_aaaa"] = aaaa
            data["assinatura_dia"] = str(contrato.data.day)
            data["assinatura_mes"] = mmm
            data["assinatura_ano"] = str(contrato.data.year)
        else:
            data["data_dd"] = "DD"
            data["data_mmm"] = "MMM"
            data["data_aaaa"] = "AAAA"
            data["assinatura_dia"] = "[dia]"
            data["assinatura_mes"] = "[mês]"
            data["assinatura_ano"] = "[ano]"

        # Compra Dates & Modality
        if compra:
            if compra.data_proposta_comercial:
                data["data_proposta"] = compra.data_proposta_comercial_dmy
            else:
                data["data_proposta"] = "[NN/NN/NNNN]"
                
            if compra.data_estimativa_orcamento:
                data["data_estimativa"] = compra.data_estimativa_orcamento_dmy
            else:
                data["data_estimativa"] = "DD/MM/AAAA"
                
            # Currency & Extenso
            if compra.valor_efetivo:
                data["valor_formatado"] = compra.valor_efetivo_brl
                data["valor_extenso"] = compra.valor_efetivo_brl_extenso
            else:
                data["valor_formatado"] = "R$ 0,00"
                data["valor_extenso"] = "zero reais"
        else:
            data["data_proposta"] = "[NN/NN/NNNN]"
            data["data_estimativa"] = "DD/MM/AAAA"
            data["valor_formatado"] = "R$ 0,00"
            data["valor_extenso"] = "zero reais"

        # Garantia
        if contrato.modalidade_garantia:
            data["garantia_modalidade"] = contrato.get_modalidade_garantia_display()
        else:
            data["garantia_modalidade"] = "[não exigida]"
            
        if contrato.valor_garantia:
            data["garantia_valor"] = contrato.valor_garantia_brl
        else:
            data["garantia_valor"] = "R$ 0,00"

        # 4. Load DOCX
        doc = Document(docx_file)
        
        # 5. Process sections with context state machine
        context_state = 'PREAMBULO'
        
        # Define a helper function to process a list of paragraphs with index
        def process_paragraphs_with_index(paragraphs):
            nonlocal context_state
            for idx, p in enumerate(paragraphs):
                if not p.text.strip():
                    continue
                # Update context state based on section headings
                context_state = cls.localizar_secao(p.text, context_state)
                
                # Check if modality specific replacement is needed
                if compra:
                    modalidade_upper = (compra.modalidade or "").upper()
                    if 'DISPENSA' in modalidade_upper:
                        if "NN/AAAA" in p.text and ("Aviso" in p.text or "Contratação Direta" in p.text):
                            cls.substituir_texto(p, "NN/AAAA", compra.numero_comprasgov or "")
                            cls.substituir_texto(p, "[SIGLA DA UNIDADE]", data.get("contratante_nome_fantasia", ""))
                    elif 'PREGÃO' in modalidade_upper or 'PREGAO' in modalidade_upper:
                        if "NN/AAAA" in p.text and ("Edital" in p.text or "Pregão" in p.text):
                            cls.substituir_texto(p, "NN/AAAA", compra.numero_comprasgov or "")
                            cls.substituir_texto(p, "[SIGLA DA UNIDADE]", data.get("contratante_nome_fantasia", ""))
                
                
                
                cls.preencher_labels_contextuais(p, context_state, data)
                # Apply general placeholders
                cls.processar_paragrafo_geral(p, data)

        # Process main document body
        process_paragraphs_with_index(doc.paragraphs)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Table cell can change or inherit context state
                    # We look at cell text to see if context changes locally for the cell
                    cell_state = cls.localizar_secao(cell.text, context_state)
                    
                    # For each paragraph in cell
                    for p_idx, p in enumerate(cell.paragraphs):
                        # Se encontramos "Representante legal", procura underscores no parágrafo anterior
                        
                        # Modalidade rules inside cells
                        if compra:
                            modalidade_upper = (compra.modalidade or "").upper()
                            if 'DISPENSA' in modalidade_upper:
                                if "NN/AAAA" in p.text and ("Aviso" in p.text or "Contratação Direta" in p.text):
                                    cls.substituir_texto(p, "NN/AAAA", compra.numero_comprasgov or "")
                                    cls.substituir_texto(p, "[SIGLA DA UNIDADE]", data.get("contratante_nome_fantasia", ""))
                            elif 'PREGÃO' in modalidade_upper or 'PREGAO' in modalidade_upper:
                                if "NN/AAAA" in p.text and ("Edital" in p.text or "Pregão" in p.text):
                                    cls.substituir_texto(p, "NN/AAAA", compra.numero_comprasgov or "")
                                    cls.substituir_texto(p, "[SIGLA DA UNIDADE]", data.get("contratante_nome_fantasia", ""))
                                    
                        
                        cls.preencher_labels_contextuais(p, cell_state, data)
                        cls.processar_paragrafo_geral(p, data)

        # Process headers and footers
        for section in doc.sections:
            if section.header:
                process_paragraphs_with_index(section.header.paragraphs)
            if section.footer:
                process_paragraphs_with_index(section.footer.paragraphs)

        # 6. Save filled document to memory
        output_io = io.BytesIO()
        doc.save(output_io)
        output_io.seek(0)
        
        # 7. Generate clean filename
        clean_num = contrato.numero.replace("/", "_").replace("\\", "_")
        filename = f"Termo_Contrato_{contratante.nome_fantasia}_{clean_num}_{contratada.nome_fantasia}.docx"
        
        return output_io, filename
