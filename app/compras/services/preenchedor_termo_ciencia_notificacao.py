import io
import re
import logging
from datetime import date
from decimal import Decimal
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from compras.models import Contrato, VinculoOrganizacao, Organizacao, Compra, Empenho



logger = logging.getLogger(__name__)


class PreenchedorTermoCienciaNotificacaoService():
    """
    Preenche o documento 'Termo de Ciência e Notificação' a partir dos dados do contrato.
    """

    @classmethod
    def substituir_texto(cls, paragraph, target, replacement):
        """
        Substitute target text with replacement inside a python-docx paragraph,
        preserving formatting (runs) as much as possible.
        Reutiliza a mesma implementação de PreenchedorContratoService.
        """
        if not target or target not in paragraph.text:
            return False

        if target == replacement:
            return False

        # Try simple run replacement
        replaced = False
        for run in paragraph.runs:
            if target in run.text:
                run.text = run.text.replace(target, replacement)
                replaced = True

        if replaced:
            return True

        # Reconstruct runs for split targets
        text = paragraph.text
        idx = text.find(target)
        max_loops = text.count(target)
        loop_count = 0

        while idx != -1 and loop_count < max_loops:
            loop_count += 1
            start_run_idx = -1
            end_run_idx = -1
            current_len = 0

            for i, run in enumerate(paragraph.runs):
                run_len = len(run.text)
                if start_run_idx == -1 and current_len + run_len > idx:
                    start_run_idx = i
                if end_run_idx == -1 and current_len + run_len >= idx + len(target):
                    end_run_idx = i
                    break
                current_len += run_len

            if start_run_idx != -1 and end_run_idx != -1:
                combined_text = ""
                for i in range(start_run_idx, end_run_idx + 1):
                    combined_text += paragraph.runs[i].text

                local_idx = combined_text.find(target)
                if local_idx != -1:
                    new_combined = combined_text[:local_idx] + replacement + combined_text[local_idx + len(target):]
                    paragraph.runs[start_run_idx].text = new_combined
                    for i in range(start_run_idx + 1, end_run_idx + 1):
                        paragraph.runs[i].text = ""
            else:
                break

            text = paragraph.text
            idx = text.find(target)

        return True

    @classmethod
    def preencher_campo_simples(cls, paragraph, label, valor):
        """
        Localiza o label no parágrafo e insere o valor após ele.
        Exemplo: 'CONTRATANTE:' → 'CONTRATANTE: Universidade de São Paulo'
        """
        if not valor:
            return

        # Verificar se o label existe no parágrafo
        if label not in paragraph.text:
            return

        # Substituir "LABEL:" por "LABEL: valor"
        cls.substituir_texto(paragraph, f"{label}:", f"{label}: {valor}")

    @classmethod
    def preencher_secao_responsavel(cls, paragraphs, start_idx, dados):
        if not dados:
            return

        campos_preenchidos = 0

        for i in range(start_idx + 1, min(start_idx + 10, len(paragraphs))):
            text = paragraphs[i].text.strip()

            if text.startswith('Nome:'):
                cls.substituir_texto(
                    paragraphs[i],
                    'Nome:',
                    f"Nome: {dados.get('nome', '')}"
                )
                campos_preenchidos += 1

            elif text.startswith('Cargo'):
                cls.substituir_texto(
                    paragraphs[i],
                    'Cargo',
                    f"Cargo: {dados.get('cargo', '')}"
                )
                campos_preenchidos += 1

            elif text.startswith('CPF:'):
                cls.substituir_texto(
                    paragraphs[i],
                    'CPF:',
                    f"CPF: {dados.get('cpf', '')}"
                )
                campos_preenchidos += 1

            # Já preencheu Nome, Cargo e CPF
            if campos_preenchidos >= 3:
                break

    @classmethod
    def localizar_secao(cls, paragraphs, titulo_secao):
        """
        Localiza o índice do parágrafo que contém o título da seção.
        """
        titulo_upper = titulo_secao.upper()
        for i, p in enumerate(paragraphs):
            if titulo_upper in p.text.upper():
                return i
        return None

    @classmethod
    def formatar_cpf(cls, cpf):
        """
        Formata CPF no padrão XXX.XXX.XXX-XX.
        """
        if not cpf:
            return ""

        # Remover caracteres não numéricos
        cpf_numeros = re.sub(r'\D', '', cpf)

        if len(cpf_numeros) != 11:
            return cpf

        return f"{cpf_numeros[:3]}.{cpf_numeros[3:6]}.{cpf_numeros[6:9]}-{cpf_numeros[9:]}"

    @classmethod
    def fill_docx(cls, docx_file, contrato):
        """
        Método principal que preenche o documento 'Termo de Ciência e Notificação'.
        
        Args:
            docx_file: Arquivo DOCX do template
            contrato: Instância do modelo Contrato
            
        Returns:
            Tuple (io.BytesIO, str): Buffer do arquivo preenchido e nome do arquivo
        """

        # 1. Buscar objetos relacionados
        compra = contrato.compra
        contratante = contrato.contratante
        contratada = contrato.contratada

        # 2. Buscar representantes
        v_contratante = VinculoOrganizacao.objects.filter(
            organizacao=contratante
        ).select_related('pessoa').first()

        if contrato.representante_contratada:
            v_contratada = contrato.representante_contratada
        else:
            v_contratada = VinculoOrganizacao.objects.filter(
                organizacao=contratada
            ).select_related('pessoa').first()

        # 3. Montar dicionário de dados
        data = {
            "contratante_nome": contratante.nome,
            "contratada_nome": contratada.nome,
            "contrato_numero": contrato.numero,
            "objeto": compra.objeto if compra else "",
            "contratante_cidade": contratante.cidade or "",
        }

        # Dados de responsáveis
        if v_contratante:
            data["responsavel_contratante"] = {
                "nome": v_contratante.pessoa.nome,
                "cargo": v_contratante.cargo or "",
                "cpf": cls.formatar_cpf(v_contratante.pessoa.cpf),
            }
        else:
            data["responsavel_contratante"] = {
                "nome": "",
                "cargo": "",
                "cpf": "",
            }

        if v_contratada:
            data["responsavel_contratada"] = {
                "nome": v_contratada.pessoa.nome,
                "cargo": v_contratada.cargo or "",
                "cpf": cls.formatar_cpf(v_contratada.pessoa.cpf),
            }
        else:
            data["responsavel_contratada"] = {
                "nome": "",
                "cargo": "",
                "cpf": "",
            }

        # Data por extenso
        if contrato.data:
            data["local_data"] = f"{data['contratante_cidade']}, {contrato.data_extenso}"
        else:
            data["local_data"] = data['contratante_cidade']

        # 4. Carregar DOCX
        doc = Document(docx_file)

        # 5. Processar parágrafos
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()

            # Substituir título
            if "ANEXO VI" in text and "TERMO DE CIÊNCIA E NOTIFICAÇÃO" in text:
                novo_texto = text.replace("ANEXO VI –  ", "").replace("ANEXO VI – ", "")
                if paragraph.runs:
                    paragraph.runs[0].text = novo_texto
                    for run in paragraph.runs[1:]:
                        run.text = ""
                continue

            # Preencher campos simples
            if text.startswith("CONTRATANTE:"):
                cls.preencher_campo_simples(paragraph, "CONTRATANTE", data["contratante_nome"])
            elif text.startswith("CONTRATADO:"):
                cls.preencher_campo_simples(paragraph, "CONTRATADO", data["contratada_nome"])
            elif text.startswith("CONTRATO Nº (DE ORIGEM):"):
                cls.preencher_campo_simples(paragraph, "CONTRATO Nº (DE ORIGEM)", data["contrato_numero"])
            elif text.startswith("OBJETO:"):
                cls.preencher_campo_simples(paragraph, "OBJETO", data["objeto"])
            elif text.startswith("LOCAL e DATA:"):
                cls.preencher_campo_simples(paragraph, "LOCAL e DATA", data["local_data"])

            # Preencher seções de responsáveis
            # RESPONSÁVEIS PELA HOMOLOGAÇÃO
            elif "RESPONSÁVEIS PELA HOMOLOGAÇÃO DO CERTAME" in text.upper():
                cls.preencher_secao_responsavel(
                    doc.paragraphs,
                    i,
                    data["responsavel_contratante"]
                )

            # RESPONSÁVEIS QUE ASSINARAM O AJUSTE
            elif text.strip().upper() == "PELA CONTRATANTE:":
                cls.preencher_secao_responsavel(
                    doc.paragraphs,
                    i,
                    data["responsavel_contratante"]
                )

            elif text.strip().upper() == "PELA CONTRATADA:":
                cls.preencher_secao_responsavel(
                    doc.paragraphs,
                    i,
                    data["responsavel_contratada"]
                )

            # ORDENADOR DE DESPESAS
            elif "ORDENADOR DE DESPESAS DA CONTRATANTE" in text.upper():
                cls.preencher_secao_responsavel(
                    doc.paragraphs,
                    i,
                    data["responsavel_contratante"]
                )

        # 6. Processar tabelas (se houver)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for j, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text.strip()

                        if text.startswith("CONTRATANTE:"):
                            cls.preencher_campo_simples(paragraph, "CONTRATANTE", data["contratante_nome"])
                        elif text.startswith("CONTRATADO:"):
                            cls.preencher_campo_simples(paragraph, "CONTRATADO", data["contratada_nome"])
                        elif text.startswith("CONTRATO Nº (DE ORIGEM):"):
                            cls.preencher_campo_simples(paragraph, "CONTRATO Nº (DE ORIGEM)", data["contrato_numero"])
                        elif text.startswith("OBJETO:"):
                            cls.preencher_campo_simples(paragraph, "OBJETO", data["objeto"])
                        elif text.startswith("LOCAL e DATA:"):
                            cls.preencher_campo_simples(paragraph, "LOCAL e DATA", data["local_data"])

        # 7. Salvar documento preenchido em memória
        output_io = io.BytesIO()
        doc.save(output_io)
        output_io.seek(0)

        # 8. Gerar nome do arquivo limpo
        clean_num = contrato.numero.replace("/", "_").replace("\\", "_")
        filename = f"Termo_Ciencia_Notificacao_{contratante.nome_fantasia.upper()}_{clean_num}_{contratada.nome_or_nome_fantasia.upper()}.docx"
        

        return output_io, filename
