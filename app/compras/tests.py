from django.test import TestCase
from decimal import Decimal
from docx import Document
from .services.preenchedor_contrato import PreenchedorContratoService

class PreenchedorContratoServiceTest(TestCase):
    
    def test_formatar_moeda_brasileira(self):
        service = PreenchedorContratoService
        self.assertEqual(service.formatar_moeda_brasileira(1234.56), "R$ 1.234,56")
        self.assertEqual(service.formatar_moeda_brasileira(0), "R$ 0,00")
        self.assertEqual(service.formatar_moeda_brasileira(None), "R$ 0,00")
        self.assertEqual(service.formatar_moeda_brasileira(1000000), "R$ 1.000.000,00")

    def test_get_month_name(self):
        service = PreenchedorContratoService
        self.assertEqual(service.get_month_name(1), "janeiro")
        self.assertEqual(service.get_month_name(5), "maio")
        self.assertEqual(service.get_month_name(12), "dezembro")
        self.assertEqual(service.get_month_name(13), "")

    def test_valor_por_extenso(self):
        service = PreenchedorContratoService
        self.assertEqual(service.valor_por_extenso(Decimal("0.00")), "zero reais")
        self.assertEqual(service.valor_por_extenso(Decimal("1.00")), "um real")
        self.assertEqual(service.valor_por_extenso(Decimal("2.00")), "dois reais")
        self.assertEqual(service.valor_por_extenso(Decimal("1.50")), "um real e cinquenta centavos")
        self.assertEqual(service.valor_por_extenso(Decimal("0.05")), "cinco centavos")
        self.assertEqual(service.valor_por_extenso(Decimal("1000.00")), "mil reais")
        self.assertEqual(service.valor_por_extenso(Decimal("1234.56")), "mil, duzentos e trinta e quatro reais e cinquenta e seis centavos")
        self.assertEqual(service.valor_por_extenso(Decimal("1000000.00")), "um milhão de reais")
        self.assertEqual(service.valor_por_extenso(Decimal("2500300.40")), "dois milhões, quinhentos mil e trezentos reais e quarenta centavos")

    def test_localizar_secao(self):
        service = PreenchedorContratoService
        self.assertEqual(service.localizar_secao("REPRESENTANTE LEGAL DO CONTRATANTE:", "PREAMBULO"), "REPRESENTANTE_CONTRATANTE")
        self.assertEqual(service.localizar_secao("Seção de Dados da Contratada", "PREAMBULO"), "CONTRATADA")
        self.assertEqual(service.localizar_secao("Qualquer outro texto", "PREAMBULO"), "PREAMBULO")

    def test_substituir_texto_in_runs(self):
        service = PreenchedorContratoService
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Texto com ")
        p.add_run("[PLACEHOLDER]")
        p.add_run(" aqui.")
        
        # Test replace when fully in a run
        res = service.substituir_texto(p, "[PLACEHOLDER]", "VALOR")
        self.assertTrue(res)
        self.assertIn("Texto com VALOR aqui.", p.text)
        
        # Test when target is split across runs
        p2 = doc.add_paragraph()
        p2.add_run("Texto com [PLACE")
        p2.add_run("HOLDER] aqui.")
        res2 = service.substituir_texto(p2, "[PLACEHOLDER]", "VALOR_DIVIDIDO")
        self.assertTrue(res2)
        self.assertIn("Texto com VALOR_DIVIDIDO aqui.", p2.text)

    def test_substituir_texto_evita_loop_infinito(self):
        service = PreenchedorContratoService
        doc = Document()
        
        # Scenario 1: target == replacement
        p = doc.add_paragraph()
        p.add_run("Texto com [PLACEHOLDER]")
        res = service.substituir_texto(p, "[PLACEHOLDER]", "[PLACEHOLDER]")
        self.assertFalse(res)
        
        # Scenario 2: target in replacement
        p2 = doc.add_paragraph()
        p2.add_run("Texto com [PLACE")
        p2.add_run("HOLDER]")
        res2 = service.substituir_texto(p2, "[PLACEHOLDER]", "Novo [PLACEHOLDER] texto")
        self.assertTrue(res2)
        self.assertIn("Texto com Novo [PLACEHOLDER] texto", p2.text)


from .models import Compra, Contrato, Organizacao
from .forms import ContratoForm

class ContratoGarantiaTest(TestCase):
    def setUp(self):
        self.organizacao_contratante = Organizacao.objects.create(
            nome="Instituição Contratante",
            cnpj="11.111.111/0001-11",
            is_propria_instituicao=True
        )
        self.organizacao_contratada = Organizacao.objects.create(
            nome="Fornecedor Contratado",
            cnpj="22.222.222/0001-22",
            is_propria_instituicao=False
        )
        self.compra = Compra.objects.create(
            numero_compra="123456789012",
            objeto="Compra Teste",
            valor_efetivo=Decimal("10000.00"),
            data_estimativa_orcamento="2026-06-01",
            data_proposta_comercial="2026-06-02"
        )

    def test_contrato_calcula_garantia_no_model(self):
        contrato = Contrato.objects.create(
            numero="CON-001/2026",
            compra=self.compra,
            contratante=self.organizacao_contratante,
            contratada=self.organizacao_contratada,
            porcentual_garantia=Decimal("5.00")
        )
        self.assertEqual(contrato.valor_garantia, Decimal("500.00"))

    def test_contrato_form_salva_campos_compra_e_garantia(self):
        contrato = Contrato.objects.create(
            numero="CON-002/2026",
            compra=self.compra,
            contratante=self.organizacao_contratante,
            contratada=self.organizacao_contratada,
        )
        
        # Simula submissão do form com edições
        form_data = {
            'numero': 'CON-002/2026-EDITADO',
            'compra': self.compra.id,
            'contratante': self.organizacao_contratante.id,
            'contratada': self.organizacao_contratada.id,
            'porcentual_garantia': '10.00',
            'valor_garantia': '',  # deve ser calculado
            'valor_efetivo': '20000.00',  # Novo valor efetivo
            'data_estimativa_orcamento': '2026-07-01',
            'data_proposta_comercial': '2026-07-02',
            'data': '2026-06-15'
        }
        
        form = ContratoForm(data=form_data, instance=contrato)
        self.assertTrue(form.is_valid())
        saved_contrato = form.save()
        
        # Verifica se o valor_garantia foi recalculado para 2000 (10% de 20000)
        self.assertEqual(saved_contrato.valor_garantia, Decimal("2000.00"))
        
        # Verifica se os campos da compra foram salvos
        self.compra.refresh_from_db()
        self.assertEqual(self.compra.valor_efetivo, Decimal("20000.00"))
        self.assertEqual(str(self.compra.data_estimativa_orcamento), '2026-07-01')
        self.assertEqual(str(self.compra.data_proposta_comercial), '2026-07-02')
